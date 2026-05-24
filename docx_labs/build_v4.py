"""
build_v4.py
===========
Correcciones respecto a v3:

1. FUENTE DE TEXTO: solo del legacy (VII.docx). Los RAW se usan ÚNICAMENTE para
   sus párrafos con fórmula OMML nativa (y el texto que llevan inline esos párrafos,
   incluyendo referencias cruzadas como "(Equation 3)").

2. SIN DUPLICADOS: cuando el legacy tiene un párrafo VML-formula, se sustituye
   por el párrafo RAW equivalente. El párrafo legacy se descarta completamente
   (texto + VML). No se añade texto de legacy adicional.

3. STRIP COMPLETO de pPr: eliminar shd, rPr, numPr de RAW (no existe en dest),
   spacing, jc, ind, etc. Solo queda pStyle.

4. RAW numPr (listas de definición de variables) → mapear a L.ListaNum1 o
   N.Viñeta1 según contexto; eliminar numPr del pPr (el estilo lo gestiona).

5. Track-changes del legacy: aceptados (ins=keep, del=remove, rPrChange=remove).
"""

import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

DIR = "/Volumes/macmini_agent/external_macos/Users/hcf/Developer/codegpt/docx_labs/"

W  = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M  = "http://schemas.openxmlformats.org/officeDocument/2006/math"

def wtag(n): return '{%s}%s' % (W, n)
def mtag(n): return '{%s}%s' % (M, n)

# ── Style ID map legacy → dest ───────────────────────────────────────────────
STYLE_MAP = {
    "Heading 4":        "ETtulo5",
    "Heading 5":        "FTtulo6",
    "Heading 6":        "GTtulo7",
    "Normal":           "ITexto",
    "Body Text":        "Textoindependiente",
    "First Paragraph":  "ITexto",
    "D.Texto":          "ITexto",
    "Texto Indep9858":  "ITexto",
    "Footer":           "ITexto",
    "F.Bullets":        "NVieta1",
    "G.Números":        "LListaNum1",
    "List Paragraph":   "NVieta1",
    "CM4+2":            "NVieta1",
    "-guiones9858":     "NVieta1",
    "Compact":          "ITexto",
    "Caption":          "ITexto",
}
# RAW styles → dest
RAW_STYLE_MAP = {
    "Body Text":        "Textoindependiente",  # standalone formula
    "First Paragraph":  "ITexto",              # inline formula + text
    "Normal":           "LListaNum1",          # list-with-formula (def variables)
    "Compact":          "LListaNum1",          # compact list items
    "Title":            "ETtulo5",
    "Heading 1":        "ETtulo5",
    "Heading 2":        "FTtulo6",
    "Heading 4":        "FTtulo6",
}

# pPr children to remove completely
STRIP_PPR_TAGS = {
    wtag('spacing'), wtag('jc'), wtag('ind'), wtag('contextualSpacing'),
    wtag('keepLines'), wtag('keepNext'), wtag('pageBreakBefore'),
    wtag('outlineLvl'), wtag('suppressAutoHyphens'), wtag('widowControl'),
    wtag('textAlignment'), wtag('adjustRightInd'), wtag('snapToGrid'),
    wtag('autoSpaceDE'), wtag('autoSpaceDN'), wtag('overflowPunct'),
    wtag('shd'),      # background shading (RAW code-block remnant)
    wtag('numPr'),    # list numbering (RAW internal numIds don't exist in dest)
    wtag('rPr'),      # paragraph-mark rPr override — drop entirely, let style rule
    wtag('suppressLineNumbers'),
}

# run rPr children to remove
STRIP_RPR_TAGS = {
    wtag('sz'), wtag('szCs'), wtag('b'), wtag('bCs'), wtag('i'), wtag('iCs'),
    wtag('color'), wtag('highlight'), wtag('noProof'), wtag('vanish'),
    wtag('strike'), wtag('dstrike'), wtag('u'), wtag('vertAlign'),
    wtag('spacing'), wtag('kern'), wtag('position'), wtag('shd'),
    wtag('fitText'), wtag('em'), wtag('effect'),
    wtag('rFonts'),   # font overrides — let style handle fonts
}
# Keep in run rPr: lang (spellcheck), eastAsianLayout, etc.
RPR_KEEP = {wtag('lang')}

# Track-change tags to remove entirely
TC_REMOVE = {
    wtag('del'), wtag('rPrChange'), wtag('pPrChange'),
    wtag('sectPrChange'), wtag('tblPrChange'), wtag('trPrChange'),
    wtag('tcPrChange'), wtag('moveFrom'), wtag('moveFromRangeStart'),
    wtag('moveFromRangeEnd'), wtag('bookmarkStart'), wtag('bookmarkEnd'),
    wtag('commentRangeStart'), wtag('commentRangeEnd'), wtag('proofErr'),
}
TC_UNWRAP = {wtag('ins'), wtag('moveTo')}  # keep content, remove wrapper

RSID_ATTRS = {
    qn('w:rsidR'), qn('w:rsidRPr'), qn('w:rsidDel'),
    qn('w:rsidDefault'), qn('w:rsidSect'), qn('w:rsid'),
    qn('w:rsidRDefault'),
}

# ── Helpers ──────────────────────────────────────────────────────────────────

def accept_revisions(elem):
    """In-place: accept all tracked changes (bottom-up)."""
    for el in reversed(list(elem.iter())):
        if el.tag in TC_REMOVE:
            p = el.getparent()
            if p is not None:
                p.remove(el)
        elif el.tag in TC_UNWRAP:
            p = el.getparent()
            if p is not None:
                idx = list(p).index(el)
                for child in list(el):
                    p.insert(idx, child)
                    idx += 1
                p.remove(el)

def strip_rsid(elem):
    for node in elem.iter():
        for attr in list(node.attrib):
            if attr in RSID_ATTRS:
                del node.attrib[attr]

def strip_ppr(pPr):
    """Remove all overrides from pPr. Keep only w:pStyle."""
    if pPr is None:
        return
    for child in list(pPr):
        if child.tag != wtag('pStyle'):
            pPr.remove(child)

def strip_run_rpr(rPr):
    """Remove visual overrides from run rPr. Keep only lang."""
    if rPr is None:
        return
    for child in list(rPr):
        if child.tag not in RPR_KEEP:
            rPr.remove(child)
    # If now empty, caller can remove the rPr
    return len(rPr) == 0

def clean_elem(elem, dest_style_id, from_raw=False):
    """
    Clone elem, accept revisions, apply dest_style_id, strip all format overrides.
    from_raw=True: also strip shd, numPr, rPr from pPr (RAW artifacts).
    """
    el = copy.deepcopy(elem)

    if not from_raw:
        accept_revisions(el)

    strip_rsid(el)

    # Fix pStyle
    pPr = el.find(wtag('pPr'))
    if pPr is None:
        pPr = OxmlElement('w:pPr')
        el.insert(0, pPr)

    pStyle = pPr.find(wtag('pStyle'))
    if pStyle is None:
        pStyle = OxmlElement('w:pStyle')
        pPr.insert(0, pStyle)
    pStyle.set(qn('w:val'), dest_style_id)

    # Strip all pPr overrides
    strip_ppr(pPr)

    # Strip run-level rPr in w:r runs (NOT inside m:oMath subtrees)
    OMATH_TAG = mtag('oMath')
    OMATHPARA_TAG = mtag('oMathPara')
    for run in el.findall('.//' + wtag('r')):
        # Check if inside oMath — skip those
        in_math = False
        for anc in run.iterancestors():
            if anc.tag in (OMATH_TAG, OMATHPARA_TAG):
                in_math = True
                break
        if in_math:
            continue
        rPr = run.find(wtag('rPr'))
        if rPr is not None:
            empty = strip_run_rpr(rPr)
            if empty:
                run.remove(rPr)

    return el

def has_native_omml(elem):
    return bool(elem.findall('.//{%s}oMath' % M))

def has_vml_formula(elem):
    xml = etree.tostring(elem).decode()
    return 'oMath' in xml

def get_text(p):
    return p.text.strip()

# ── Load documents ────────────────────────────────────────────────────────────
print("Loading documents...")
legacy   = Document(DIR + "VII. RIESGO DE CONTRAPARTIDA.docx")
dest     = Document(DIR + "5.1 RIESGO DE CONTRAPARTIDA - Marco teórico y metodología.docx")
raw_fx   = Document(DIR + "20240508_RAW_FX_difussion_models.docx")
raw_ir   = Document(DIR + "20240509_RAW_IR_difussion_models.docx")
raw_inf  = Document(DIR + "20240510_RAW_INF_difussion_models.docx")
raw_comm = Document(DIR + "20240510_RAW_COMM_difussion_models.docx")

dest_style_id_map = {s.name: s.style_id for s in dest.styles}

# ── Build RAW formula paragraph queues (only those with native OMML) ─────────
def get_raw_formula_queue(raw_doc):
    """Return ordered list of (elem, has_oMathPara, src_style) for OMML paras."""
    result = []
    for p in raw_doc.paragraphs:
        if has_native_omml(p._element):
            has_para = bool(p._element.findall('.//{%s}oMathPara' % M))
            result.append((p._element, has_para, p.style.name if p.style else 'Normal'))
    return result

raw_queues = {
    'FX':   get_raw_formula_queue(raw_fx),
    'IR':   get_raw_formula_queue(raw_ir),
    'INF':  get_raw_formula_queue(raw_inf),
    'COMM': get_raw_formula_queue(raw_comm),
}
for k, v in raw_queues.items():
    print(f"  RAW {k}: {len(v)} OMML formula paragraphs queued")

cursors = {k: 0 for k in raw_queues}

# ── Section ranges in legacy ──────────────────────────────────────────────────
SECTIONS = [
    ('FX',   406, 503),
    ('IR',   503, 594),
    ('INF',  594, 641),
    ('COMM', 641, 699),
]

# ── Find insertion point in dest ──────────────────────────────────────────────
anchor = None
for p in dest.paragraphs:
    if p.style.name == "C.Título3" and "REQUERIMIENTOS" in p.text.upper():
        anchor = p._element
        print(f"\nInsertion anchor: '{p.text[:60]}'")
        break
assert anchor is not None, "Anchor not found"

body = dest.element.body

def insert_before(elem):
    body.insert(list(body).index(anchor), elem)

# ── Process sections ──────────────────────────────────────────────────────────
total_inserted = total_formulas = total_skipped = 0

for sec_name, sec_start, sec_end in SECTIONS:
    print(f"\n{'='*60}\nSection {sec_name} [{sec_start}..{sec_end-1}]")
    queue  = raw_queues[sec_name]
    cursor = cursors[sec_name]
    sec_i = sec_f = 0

    leg_paras = legacy.paragraphs

    for li in range(sec_start, sec_end):
        p    = leg_paras[li]
        sty  = p.style.name if p.style else "Normal"
        text = get_text(p)
        is_vml    = has_vml_formula(p._element)
        is_native = has_native_omml(p._element)

        # ── A: Heading → fresh paragraph, zero contamination ─────────────────
        if sty in ("Heading 4", "Heading 5", "Heading 6"):
            if not text:
                total_skipped += 1
                continue
            dest_sid = STYLE_MAP[sty]
            new_p = OxmlElement('w:p')
            new_pPr = OxmlElement('w:pPr')
            ps = OxmlElement('w:pStyle')
            ps.set(qn('w:val'), dest_sid)
            new_pPr.append(ps)
            new_p.append(new_pPr)
            new_r = OxmlElement('w:r')
            new_t = OxmlElement('w:t')
            new_t.text = text
            new_r.append(new_t)
            new_p.append(new_r)
            insert_before(new_p)
            print(f"  [{li}] HEADING ({sty}→{dest_sid}): {text[:60]}")
            sec_i += 1
            continue

        # ── B: VML formula paragraph → replace with RAW OMML paragraph ───────
        if is_vml and not is_native:
            if cursor < len(queue):
                raw_elem, has_omathpara, raw_sty = queue[cursor]
                cursor += 1

                # Determine destination style for this formula paragraph
                if has_omathpara and not p.text.strip():
                    # Standalone block formula
                    dest_sid = "Textoindependiente"
                else:
                    # Inline formula: use RAW style mapping
                    dest_sid = RAW_STYLE_MAP.get(raw_sty, "ITexto")

                new_elem = clean_elem(raw_elem, dest_sid, from_raw=True)
                insert_before(new_elem)
                sec_i += 1
                sec_f += 1
                # NOTE: we do NOT also insert the legacy text —
                # the RAW paragraph already carries the inline text + formula.
            else:
                print(f"  [{li}] WARN: RAW queue exhausted for {sec_name}")
                total_skipped += 1
            continue

        # ── C: Empty paragraph without formula → skip ────────────────────────
        if not text and not is_native:
            total_skipped += 1
            continue

        # ── D: Normal text paragraph from legacy ─────────────────────────────
        dest_sid = STYLE_MAP.get(sty, "ITexto")
        new_elem = clean_elem(p._element, dest_sid, from_raw=False)
        # Remove any residual pict/VML wrappers
        for pict in new_elem.findall('.//' + wtag('pict')):
            pp = pict.getparent()
            if pp is not None:
                pp.remove(pict)
        # Skip if after cleaning there's no content
        texts = [t.text for t in new_elem.findall('.//' + wtag('t')) if t.text]
        if not texts and not has_native_omml(new_elem):
            total_skipped += 1
            continue
        insert_before(new_elem)
        sec_i += 1

    cursors[sec_name] = cursor
    total_inserted += sec_i
    total_formulas += sec_f
    print(f"  → {sec_i} paragraphs inserted, {sec_f} formulas from RAW")

print(f"\n{'='*60}")
print(f"TOTAL inserted : {total_inserted}")
print(f"TOTAL formulas : {total_formulas}")
print(f"TOTAL skipped  : {total_skipped}")

# ── Clean document settings ───────────────────────────────────────────────────
settings = dest.settings.element
for tc in settings.findall('.//{%s}trackChanges' % W):
    settings.remove(tc)

# ── Save ─────────────────────────────────────────────────────────────────────
out = DIR + "5.1 RIESGO DE CONTRAPARTIDA - Marco teórico y metodología_v4.docx"
dest.save(out)
print(f"\nSaved: {out}")

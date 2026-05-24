"""
Script para insertar las secciones FX, IR, Inflación y Commodities
del documento legacy VII.docx en el documento destino 5.1.docx,
aplicando los estilos corporativos del dotx.

Estrategia:
- Extraer párrafos [406..698] del legacy (FX, IR, Inflación, Commodities)
- Clonar cada párrafo con su XML (preservando fórmulas OMML)
- Cambiar el nombre de estilo al equivalente en la plantilla
- Insertar antes del heading REQUERIMIENTOS DE CAPITAL (párrafo [229] del destino,
  que es el primer C.Título3 después de MODELIZACIÓN MATEMÁTICA)
"""

import copy
from docx import Document
from docx.oxml.ns import qn
from lxml import etree

DIR = "/Volumes/macmini_agent/external_macos/Users/hcf/Developer/codegpt/docx_labs/"

# ── Mapeo de estilos legacy → destino ──────────────────────────────────────
STYLE_MAP = {
    # Headings
    "Heading 4":          "E.Título5",
    "Heading 5":          "F.Título6",
    "Heading 6":          "G.Título7",
    # Cuerpo de texto
    "Normal":             "I.Texto",
    "Body Text":          "Body Text",   # mantener para bloques de fórmula standalone
    "First Paragraph":    "I.Texto",
    "D.Texto":            "I.Texto",
    "Texto Indep9858":    "I.Texto",
    # Listas
    "F.Bullets":          "N.Viñeta1",
    "G.Números":          "L.ListaNum1",
    "List Paragraph":     "N.Viñeta1",
    "CM4+2":              "N.Viñeta1",
    "-guiones9858":       "N.Viñeta1",
    # Footer (aparece erráticamente en legacy como texto normal)
    "Footer":             "I.Texto",
}

def map_style(style_name):
    return STYLE_MAP.get(style_name, "I.Texto")

def has_omml(para_elem):
    """True si el párrafo contiene una fórmula OMML."""
    return para_elem.find('.//' + qn('m:oMath')) is not None

def clone_para_with_style(src_para, dest_doc, dest_style_id_map):
    """
    Clona el elemento XML del párrafo fuente, cambia el nombre del estilo
    al equivalente del destino y devuelve el elemento listo para insertar.
    dest_style_id_map: dict {style_name -> style_id}
    """
    elem = copy.deepcopy(src_para._element)

    # Obtener nombre del estilo original
    src_style = src_para.style.name if src_para.style else "Normal"
    new_style_name = map_style(src_style)

    # Obtener el style_id real del destino (Word usa IDs, no nombres, en w:pStyle)
    if new_style_name not in dest_style_id_map:
        new_style_name = "I.Texto"
    new_style_id = dest_style_id_map[new_style_name]

    # Modificar el pPr/pStyle en el XML clonado
    pPr = elem.find(qn('w:pPr'))
    if pPr is None:
        pPr = etree.SubElement(elem, qn('w:pPr'))
        elem.insert(0, pPr)

    pStyle = pPr.find(qn('w:pStyle'))
    if pStyle is None:
        pStyle = etree.SubElement(pPr, qn('w:pStyle'))
        pPr.insert(0, pStyle)
    pStyle.set(qn('w:val'), new_style_id)

    # Eliminar resaltado (highlight) si existe en rPr de párrafo
    for rPr in elem.findall('.//' + qn('w:rPr')):
        highlight = rPr.find(qn('w:highlight'))
        if highlight is not None:
            rPr.remove(highlight)
        # Eliminar color explícito heredado del legacy si está hardcodeado
        # (dejar que el estilo del destino imponga el color)
        color = rPr.find(qn('w:color'))
        if color is not None:
            val = color.get(qn('w:val'), '')
            if val.upper() not in ('4ABDF0', '129BD8', '0C6790', '595959', '5A5B5D', 'AUTO'):
                rPr.remove(color)

    return elem


# ── Cargar documentos ────────────────────────────────────────────────────────
legacy = Document(DIR + "VII. RIESGO DE CONTRAPARTIDA.docx")
dest   = Document(DIR + "5.1 RIESGO DE CONTRAPARTIDA - Marco teórico y metodología.docx")

# Build style_id map for destination
dest_style_id_map = {s.name: s.style_id for s in dest.styles}

legacy_paras = legacy.paragraphs
dest_paras   = dest.paragraphs

# ── Identificar rango de extracción en legacy ────────────────────────────────
# Desde "Instrumentos de Tipo de Cambio" (H4, idx ~406) hasta antes de
# "CALIBRACIÓN DE LOS PARÁMETROS" (H3, idx ~699)
EXTRACT_START_TEXT = "Instrumentos de Tipo de Cambio"
EXTRACT_END_TEXT   = "CALIBRACIÓN DE LOS PARÁMETROS"

start_idx = None
end_idx   = None

for i, p in enumerate(legacy_paras):
    if start_idx is None and "Instrumentos de Tipo de Cambio" in p.text and "Heading 4" in p.style.name:
        start_idx = i
        print(f"Extract START: [{i}] {p.style.name} | {p.text[:60]}")
    if start_idx is not None and end_idx is None and "CALIBRACIÓN" in p.text.upper() and "Heading 3" in p.style.name:
        end_idx = i
        print(f"Extract END:   [{i}] {p.style.name} | {p.text[:60]}")
        break

if start_idx is None or end_idx is None:
    raise ValueError(f"No se encontraron marcadores: start={start_idx}, end={end_idx}")

source_paras = legacy_paras[start_idx:end_idx]
print(f"\nExtrayendo {len(source_paras)} párrafos del legacy [{start_idx}..{end_idx-1}]")

# ── Identificar punto de inserción en destino ────────────────────────────────
# Insertar ANTES del heading C.Título3 'REQUERIMIENTOS DE CAPITAL...'
insert_before_idx = None
for i, p in enumerate(dest_paras):
    if p.style.name == "C.Título3" and "REQUERIMIENTOS" in p.text.upper():
        insert_before_idx = i
        print(f"\nInsertion point: [{i}] {p.style.name} | {p.text[:60]}")
        break

if insert_before_idx is None:
    raise ValueError("No se encontró el heading REQUERIMIENTOS DE CAPITAL en el destino")

# El párrafo de referencia (anchor) en el body XML del destino
anchor_para = dest_paras[insert_before_idx]
anchor_elem = anchor_para._element
body = dest.element.body

# ── Insertar párrafos clonados ───────────────────────────────────────────────
inserted = 0
skipped  = 0

for src_para in source_paras:
    # Saltar párrafos vacíos sin fórmula (reducir ruido)
    text = src_para.text.strip()
    style = src_para.style.name if src_para.style else "Normal"
    is_formula_only = has_omml(src_para._element) and not text
    is_empty = not text and not has_omml(src_para._element)

    if is_empty and style not in ("Heading 4", "Heading 5", "Heading 6", "Body Text"):
        skipped += 1
        continue  # saltar vacíos sin fórmula

    new_elem = clone_para_with_style(src_para, dest, dest_style_id_map)
    body.insert(list(body).index(anchor_elem), new_elem)
    inserted += 1

print(f"\nInsertados: {inserted} párrafos")
print(f"Omitidos (vacíos): {skipped} párrafos")

# ── Guardar resultado ────────────────────────────────────────────────────────
out_path = DIR + "5.1 RIESGO DE CONTRAPARTIDA - Marco teórico y metodología_v2.docx"
dest.save(out_path)
print(f"\nGuardado: {out_path}")
